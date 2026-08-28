"""Generate publication artifacts for the KidSpark AI handoff package.

Markdown remains the maintainable source. This tool exports diagrams, DOCX, PDF,
and OpenAPI snapshots without reading runtime secrets.
"""

from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
HANDOFF = DOCS / "handoff"
IMAGES = DOCS / "images" / "handoff"
OPENAPI = DOCS / "openapi"
RUNTIME_PYTHON = Path(sys.executable)

NAVY = "#17324D"
TEAL = "#2F7F8D"
SKY = "#E8F4F7"
CORAL = "#EF5B54"
GOLD = "#E8B84C"
INK = "#243142"
MUTED = "#667085"
LIGHT = "#F6F8FA"
GREEN = "#3A7D44"


@dataclass(frozen=True)
class DiagramNode:
    key: str
    title: str
    items: list[str]
    kind: str = "service"


@dataclass(frozen=True)
class DiagramSection:
    title: str
    subtitle: str
    nodes: list[DiagramNode]


@dataclass(frozen=True)
class DiagramEdge:
    source: str
    target: str
    label: str
    dashed: bool = False


@dataclass(frozen=True)
class Diagram:
    slug: str
    title: str
    purpose: str
    sections: list[DiagramSection]
    edges: list[DiagramEdge]


DIAGRAMS = [
    Diagram(
        "system-context",
        "C4 Level 1 - System Context",
        "Who uses KidSpark AI, what value crosses the system boundary, and which external platforms it depends on.",
        [
            DiagramSection("People and classroom", "Primary users and beneficiaries", [
                DiagramNode("teacher", "Teacher", ["Uploads or types a story", "Plans and approves each checkpoint", "Downloads the lesson bundle"], "actor"),
                DiagramNode("students", "Students and classroom", ["Read, discuss, build, test", "Use the activity guide and slides", "Assemble validated BrickSmart model"], "actor"),
                DiagramNode("team", "Jacobs research and operations team", ["Maintains curriculum corpus", "Operates deployment and ingestion", "Reviews outcomes and limitations"], "actor"),
            ]),
            DiagramSection("System of interest", "KidSpark AI / BrickSmart", [
                DiagramNode("kidspark", "KidSpark AI", ["Teacher-guided six-step workflow", "Evidence-grounded lesson coaching", "3D-to-physical build planning", "Three coordinated classroom PDFs", "Human approval before consequential stages"], "system"),
            ]),
            DiagramSection("External systems", "Managed platforms and specialist services", [
                DiagramNode("vertex", "Google Vertex AI", ["Gemini generation", "Text embeddings", "Primary and fallback model routing"], "external"),
                DiagramNode("data", "GCP data platform", ["Cloud SQL PostgreSQL + pgvector", "GCS raw and processed corpus", "Secret Manager and Cloud Logging"], "datastore"),
                DiagramNode("hyper3d", "Hyper3D platform", ["Rodin text-to-3D generation", "Bang semantic segmentation", "Asynchronous task polling"], "external"),
            ]),
        ],
        [
            DiagramEdge("teacher", "kidspark", "story, decisions, approvals"),
            DiagramEdge("kidspark", "teacher", "previews, evidence, documents", True),
            DiagramEdge("kidspark", "students", "classroom lesson and build"),
            DiagramEdge("team", "kidspark", "corpus and operations"),
            DiagramEdge("kidspark", "vertex", "generation and embeddings"),
            DiagramEdge("kidspark", "data", "SQL, objects, secrets, logs"),
            DiagramEdge("kidspark", "hyper3d", "3D jobs and artifacts"),
        ],
    ),
    Diagram(
        "application-components",
        "C4 Level 2 - Application Containers and Components",
        "The deployable application boundary and the responsibility of each major runtime component.",
        [
            DiagramSection("Presentation container", "Streamlit web application", [
                DiagramNode("ui", "Teacher workflow UI", ["Six-step rail and checkpoints", "Planning conversation", "Visual model and segment review", "Document preview and download"], "system"),
                DiagramNode("client", "API client and poller", ["Session-scoped HTTP calls", "Long-job progress polling", "Renders API state; does not own gates"], "service"),
            ]),
            DiagramSection("Application container", "FastAPI orchestration and authority", [
                DiagramNode("api", "Session API routers", ["Upload, message, confirm", "Model, segments, build, documents", "Download and runtime endpoints"], "system"),
                DiagramNode("state", "Session and readiness service", ["Planning contract", "Confirmation gates", "Artifact generation IDs", "Stale-output invalidation"], "service"),
                DiagramNode("jobs", "Background job manager", ["Threaded long-running jobs", "Progress events and status", "Normalized provider errors"], "service"),
                DiagramNode("agents", "Agent orchestration", ["Story analysis", "Planning coach and readiness guard", "Prompt and document agents"], "service"),
            ]),
            DiagramSection("Domain and adapter layer", "Focused services behind stable interfaces", [
                DiagramNode("retrieval", "Retrieval adapter", ["Grade-band normalization", "Cloud SQL or service retrieval", "Evidence trace and static fallback"], "service"),
                DiagramNode("llm", "Vertex Gemini adapter", ["Structured text/JSON generation", "Primary/fallback models", "ADC or approved key auth"], "external"),
                DiagramNode("build3d", "3D orchestration", ["Rodin and Bang clients", "Notebook physicalization", "Bounded automatic simplification"], "service"),
                DiagramNode("planner", "Validated BrickSmart runtime", ["Catalog and geometry mapping", "Inventory-aware planning", "True step renderer and validation"], "service"),
                DiagramNode("documents", "Publication service", ["Three audience-specific documents", "Content/image validation", "PDF and source artifacts"], "service"),
            ]),
            DiagramSection("State and artifacts", "Runtime persistence boundaries", [
                DiagramNode("memory", "In-memory session store", ["Teacher workflow state", "Job status and approvals", "Lost on instance replacement"], "datastore"),
                DiagramNode("work", "Ephemeral work directory", ["OBJ and segmented assets", "Notebook images and CSVs", "Generated PDFs"], "datastore"),
                DiagramNode("corpus", "Durable corpus stores", ["Cloud SQL pgvector", "GCS source and processed assets", "Not the session system of record"], "datastore"),
            ]),
        ],
        [
            DiagramEdge("ui", "client", "user actions"), DiagramEdge("client", "api", "loopback HTTP/JSON"),
            DiagramEdge("api", "state", "read/write workflow"), DiagramEdge("api", "jobs", "submit and poll"),
            DiagramEdge("state", "memory", "session state"), DiagramEdge("jobs", "work", "artifacts and progress"),
            DiagramEdge("agents", "retrieval", "grounding query"), DiagramEdge("agents", "llm", "generation"),
            DiagramEdge("jobs", "build3d", "3D stages"), DiagramEdge("build3d", "planner", "physical plan"),
            DiagramEdge("jobs", "documents", "approved bundle context"), DiagramEdge("retrieval", "corpus", "vector and policy lookup"),
        ],
    ),
    Diagram(
        "gcp-deployment",
        "GCP Deployment, Network, and Trust Boundaries",
        "How the single Cloud Run service reaches managed GCP resources and the external Hyper3D boundary.",
        [
            DiagramSection("Public client", "Internet-facing boundary", [
                DiagramNode("browser", "Teacher browser", ["HTTPS to /kidspark", "Receives Streamlit UI and downloads", "No direct access to internal API port"], "actor"),
            ]),
            DiagramSection("Cloud Run service: kidspark", "Project kidspark-499901, us-central1", [
                DiagramNode("launcher", "cloudrun_start.py", ["Starts and supervises two processes", "Propagates child failures"], "service"),
                DiagramNode("streamlit", "Streamlit process", ["0.0.0.0:$PORT (normally 8080)", "Public /kidspark route", "Calls FastAPI over loopback"], "system"),
                DiagramNode("fastapi", "FastAPI process", ["127.0.0.1:8001 only", "Session orchestration and downloads", "Health, readiness, runtime settings"], "system"),
                DiagramNode("runtime", "Runtime identity and local disk", ["Cloud Run service account", "Application Default Credentials", "Ephemeral /work and generated files", "2 CPU / 4 GiB / 3600s baseline"], "datastore"),
            ]),
            DiagramSection("Managed GCP services", "Private project resources and Google APIs", [
                DiagramNode("sql", "Cloud SQL PostgreSQL", ["pgvector extension", "document_bundle, pdf_node", "standard_rules and vector search"], "datastore"),
                DiagramNode("gcs", "Cloud Storage", ["kidspark-project-data", "kidspark-data-processed", "Corpus lineage and page assets"], "datastore"),
                DiagramNode("vertex", "Vertex AI", ["Gemini 3.6 Flash primary", "Gemini 3.5 Flash fallback", "gemini-embedding-001 (3072d)"], "external"),
                DiagramNode("secrets", "Secret Manager", ["Database URL", "Hyper3D credential", "Optional Gemini API key"], "external"),
                DiagramNode("ops", "Cloud Logging and Monitoring", ["Request/stage correlation", "Provider duration and fallback", "Failures without sensitive payloads"], "external"),
            ]),
            DiagramSection("External processor", "Third-party trust boundary", [
                DiagramNode("hyper3d", "Hyper3D API", ["Rodin generation endpoint", "Bang segmentation endpoint", "Bearer credential", "Asynchronous task IDs and downloads"], "external"),
            ]),
        ],
        [
            DiagramEdge("browser", "streamlit", "HTTPS :443"), DiagramEdge("launcher", "streamlit", "starts"),
            DiagramEdge("launcher", "fastapi", "starts"), DiagramEdge("streamlit", "fastapi", "HTTP 127.0.0.1:8001"),
            DiagramEdge("fastapi", "runtime", "state and artifacts"), DiagramEdge("fastapi", "sql", "Cloud SQL socket / psycopg"),
            DiagramEdge("fastapi", "gcs", "Google Storage API"), DiagramEdge("fastapi", "vertex", "Vertex API via ADC"),
            DiagramEdge("runtime", "secrets", "secret injection"), DiagramEdge("fastapi", "ops", "structured stdout/stderr"),
            DiagramEdge("fastapi", "hyper3d", "HTTPS REST + polling"),
        ],
    ),
    Diagram(
        "teacher-session-sequence",
        "Teacher Session Orchestration and Confirmation Gates",
        "The six teacher-visible stages, authoritative gates, long-running jobs, and invalidation paths.",
        [
            DiagramSection("Teacher and UI", "Explicit human decisions", [
                DiagramNode("s1", "1. Upload story", ["PDF or typed story", "Review extraction and framework anchors", "Confirm source"], "actor"),
                DiagramNode("s2", "2. Plan with coach", ["Iterative guided conversation", "Watch lesson-component checklist", "Confirm only when all required fields exist"], "actor"),
                DiagramNode("s3", "3. Approve model preview", ["Edit visual/build constraints", "Wait for Rodin job", "Approve model or regenerate"], "actor"),
                DiagramNode("s4", "4. Review segments/connectors", ["Inspect Bang and voxel views", "Edit labels and connector intent", "Approve or accept recovery guidance"], "actor"),
                DiagramNode("s5", "5. Approve build plan", ["Review inventory and every true step", "Approve all or individual steps"], "actor"),
                DiagramNode("s6", "6. Validate lesson bundle", ["Preview three documents", "Refine independently", "Approve and download"], "actor"),
            ]),
            DiagramSection("FastAPI state machine", "Backend authority", [
                DiagramNode("g1", "Story gate", ["Extraction non-empty", "Analysis and evidence available"], "decision"),
                DiagramNode("g2", "Planning readiness gate", ["Grade, duration, concept, goals", "Build object, movement, literacy/SEL", "Constraints and framework matches"], "decision"),
                DiagramNode("g3", "Model gate", ["Current model generation ID", "Rodin artifact complete", "Teacher confirmation"], "decision"),
                DiagramNode("g4", "Physicalization gate", ["Current segment generation ID", "Block and segment budgets", "Review-ready validated status"], "decision"),
                DiagramNode("g5", "Build gate", ["Inventory feasible", "True steps and images available", "Teacher step approvals"], "decision"),
                DiagramNode("g6", "Document gate", ["All three documents valid", "Required images and sections", "Teacher document approvals"], "decision"),
            ]),
            DiagramSection("Asynchronous work", "Polled jobs and generated evidence", [
                DiagramNode("storyjob", "Story analysis + RAG", ["Extract, analyze, retrieve", "Return evidence trace"], "service"),
                DiagramNode("rodinjob", "Rodin job", ["Submit once", "Poll provider status", "Download preview and OBJ"], "external"),
                DiagramNode("segmentjob", "Bang + notebook job", ["Segment approved OBJ", "Voxelize and auto-simplify", "Render multiview and connector tables"], "external"),
                DiagramNode("planjob", "Validated planner job", ["Map catalog and inventory", "Validate placement and sequence", "Render true build steps"], "service"),
                DiagramNode("docjob", "Document bundle job", ["Compose three audiences", "Validate content and images", "Publish PDFs and sources"], "service"),
            ]),
        ],
        [
            DiagramEdge("s1", "g1", "confirm"), DiagramEdge("g1", "storyjob", "analyze"),
            DiagramEdge("s2", "g2", "message / confirm"), DiagramEdge("g2", "s3", "ready"),
            DiagramEdge("s3", "g3", "approve"), DiagramEdge("g3", "rodinjob", "submit / poll"),
            DiagramEdge("s4", "g4", "approve"), DiagramEdge("g4", "segmentjob", "submit / poll"),
            DiagramEdge("s5", "g5", "approve steps"), DiagramEdge("g5", "planjob", "validate"),
            DiagramEdge("s6", "g6", "approve documents"), DiagramEdge("g6", "docjob", "generate / refine"),
            DiagramEdge("g4", "s3", "model regeneration required", True),
            DiagramEdge("g3", "s4", "invalidate old segments", True),
        ],
    ),
    Diagram(
        "rag-ingestion-retrieval",
        "RAG Offline Ingestion and Online Retrieval",
        "How curriculum evidence becomes traceable pgvector records and then grounds each teacher-planning turn.",
        [
            DiagramSection("Offline corpus ingestion", "Administrative/research workflow", [
                DiagramNode("sources", "Source documents", ["Framework and standards PDFs", "Teacher lesson plans", "Activity guides and slide companions", "Catalog and policy data"], "artifact"),
                DiagramNode("extract", "Docling extraction", ["Page-aware text and tables", "Document/node identity", "Image and page-crop references"], "service"),
                DiagramNode("enrich", "Normalize and enrich", ["Canonical grade_band", "Bundle and content roles", "Selective visual captions", "Provenance and checksums"], "service"),
                DiagramNode("embed", "Embedding pipeline", ["gemini-embedding-001", "3072-dimensional text vectors", "Batched, quota-aware writes", "Selective visual embeddings only"], "external"),
            ]),
            DiagramSection("Durable knowledge stores", "Shared production data", [
                DiagramNode("rawgcs", "GCS source bucket", ["Original PDFs", "Immutable source lineage"], "datastore"),
                DiagramNode("processed", "GCS processed bucket", ["Knowledge_chunks bundles", "Page crops and extracted artifacts", "Stable gs:// URIs"], "datastore"),
                DiagramNode("bundle", "document_bundle", ["Document identity and metadata", "Processing/version status", "GCS lineage"], "datastore"),
                DiagramNode("nodes", "pdf_node", ["Normalized content", "grade_band and role", "embedding vector(3072)", "Bundle/page provenance"], "datastore"),
                DiagramNode("rules", "standard_rules", ["Framework and policy anchors", "Structured grade/applicability filters"], "datastore"),
            ]),
            DiagramSection("Online teacher-turn retrieval", "Read-only classroom path", [
                DiagramNode("query", "Retrieval request", ["Teacher turn + planning state", "Canonical grade band", "Seed k and timeout"], "service"),
                DiagramNode("qembed", "Query understanding", ["Optional expansion", "Gemini query embedding", "No model instructions from corpus"], "external"),
                DiagramNode("search", "Hybrid evidence selection", ["Exact grade filter", "Vector nearest neighbors", "Policy lookup", "Deduplicate and rank"], "service"),
                DiagramNode("trace", "Evidence pack", ["Source and node IDs", "Excerpt/summary and relevance", "Retrieval mode and lineage", "Bounded prompt context"], "artifact"),
                DiagramNode("coach", "Planning coach", ["Grounded recommendations", "Deterministic readiness guard", "Static evidence fallback if degraded"], "system"),
            ]),
        ],
        [
            DiagramEdge("sources", "extract", "parse"), DiagramEdge("extract", "enrich", "nodes + assets"),
            DiagramEdge("enrich", "embed", "text/caption batches"), DiagramEdge("sources", "rawgcs", "store originals"),
            DiagramEdge("enrich", "processed", "write artifacts"), DiagramEdge("enrich", "bundle", "upsert metadata"),
            DiagramEdge("embed", "nodes", "upsert vectors"), DiagramEdge("enrich", "rules", "load policies"),
            DiagramEdge("query", "qembed", "normalize + embed"), DiagramEdge("qembed", "search", "query vector"),
            DiagramEdge("nodes", "search", "filtered vector candidates"), DiagramEdge("rules", "search", "policy anchors"),
            DiagramEdge("search", "trace", "ranked evidence"), DiagramEdge("trace", "coach", "untrusted reference context"),
        ],
    ),
    Diagram(
        "logical-data-model",
        "Logical Data and Artifact Model",
        "The main durable and ephemeral records that connect teacher intent, evidence, 3D generations, validated steps, and documents.",
        [
            DiagramSection("Teacher workflow domain", "Ephemeral session authority", [
                DiagramNode("session", "Session", ["session_id, stage, timestamps", "current generation IDs", "progress and approvals"], "datastore"),
                DiagramNode("planning", "PlanningState", ["grade, duration, concept, goals", "build object and constraints", "moving/static parts, literacy, SEL"], "datastore"),
                DiagramNode("evidencepack", "EvidenceTrace", ["document/node identity", "grade band and relevance", "retrieval mode and excerpt"], "artifact"),
            ]),
            DiagramSection("Knowledge domain", "Durable Cloud SQL and GCS corpus", [
                DiagramNode("dbbundle", "document_bundle", ["bundle identity and metadata", "source/processed GCS URIs", "processing version/status"], "datastore"),
                DiagramNode("dbnode", "pdf_node", ["page/node content and role", "grade_band", "embedding vector(3072)", "bundle foreign key"], "datastore"),
                DiagramNode("dbrule", "standard_rules", ["framework/policy statement", "grade/applicability metadata", "source lineage"], "datastore"),
            ]),
            DiagramSection("3D and physical build domain", "Generation-scoped artifacts", [
                DiagramNode("modelgen", "ModelGeneration", ["Rodin prompt and task ID", "preview/model artifact", "approval and generation number"], "artifact"),
                DiagramNode("segmentgen", "SegmentGeneration", ["Bang task and segment assets", "labels, interfaces, movement map", "voxel attempts and diagnostics"], "artifact"),
                DiagramNode("buildplan", "ValidatedBuildPlan", ["status and inventory feasibility", "catalog-mapped parts", "true instruction steps", "rendered step images"], "artifact"),
            ]),
            DiagramSection("Publication domain", "Teacher-downloadable outputs", [
                DiagramNode("bundleout", "LessonBundle", ["approved context snapshot", "generation/version state", "three document records"], "artifact"),
                DiagramNode("doc", "GeneratedDocument", ["lesson_plan | activity_guide | slide_companion", "markdown/JSON/PDF paths", "validation and teacher approval"], "artifact"),
                DiagramNode("files", "Artifact files", ["OBJ/GLB and segmented models", "PNGs, CSVs, JSON", "PDFs in ephemeral work directory"], "datastore"),
            ]),
        ],
        [
            DiagramEdge("session", "planning", "1 owns 1"), DiagramEdge("planning", "evidencepack", "references many"),
            DiagramEdge("evidencepack", "dbnode", "traces to"), DiagramEdge("dbbundle", "dbnode", "1 contains many"),
            DiagramEdge("dbrule", "evidencepack", "grounds"), DiagramEdge("session", "modelgen", "creates versions"),
            DiagramEdge("modelgen", "segmentgen", "approved model only"), DiagramEdge("segmentgen", "buildplan", "physical plan"),
            DiagramEdge("buildplan", "bundleout", "approved source"), DiagramEdge("bundleout", "doc", "contains three"),
            DiagramEdge("modelgen", "files", "writes"), DiagramEdge("segmentgen", "files", "writes"),
            DiagramEdge("buildplan", "files", "writes"), DiagramEdge("doc", "files", "writes"),
        ],
    ),
    Diagram(
        "rodin-validated-build",
        "Rodin-to-Validated BrickSmart Physicalization",
        "The geometry, semantic, connector, inventory, and teacher-review path from approved intent to true build steps.",
        [
            DiagramSection("Approved design intent", "Teacher-approved semantic contract", [
                DiagramNode("intent", "model_task_context.json", ["Object type and broad required parts", "One or more explicit movement intents", "Static parts and connector intent", "Inventory basis and build budgets"], "artifact"),
                DiagramNode("guard", "Pre-generation build guard", ["Chunky 2x2-compatible geometry", "Max semantic and moving parts", "Broad contacts; no micro-details", "Rodin prompt length and safety checks"], "decision"),
            ]),
            DiagramSection("External geometry services", "Probabilistic 3D generation", [
                DiagramNode("rodin", "Rodin generation", ["Submit visual prompt", "Poll asynchronous task", "Download preview and OBJ", "Teacher approves or regenerates"], "external"),
                DiagramNode("bang", "Bang segmentation", ["Runs only on current approved OBJ", "Returns semantic OBJ regions", "Source-segment count measured early", "Task ID tied to model generation"], "external"),
            ]),
            DiagramSection("Notebook-derived physicalization", "Deterministic geometry processing with bounded tuning", [
                DiagramNode("voxel", "Voxelize and consolidate", ["Sample occupancy at candidate sizes", "Protect moving/required regions", "Merge compatible static fragments", "Track preservation and budgets"], "service"),
                DiagramNode("contacts", "Interfaces and connectors", ["Detect adjacent segment contacts", "Map teacher movement intent", "Infer axle/pivot/static candidates", "Emit editable CSV tables"], "service"),
                DiagramNode("renders", "Review artifacts", ["Color-coded segment visualization", "Eight-view placement sheet", "Final block approximation", "Per-step isometric/multiview images"], "artifact"),
            ]),
            DiagramSection("Validated BrickSmart planner", "Catalog and finite-kit authority", [
                DiagramNode("catalog", "Catalog and inventory mapping", ["Map regions to allowed block types", "Compile standard-kit ledger", "Apply symmetry and functional rules"], "service"),
                DiagramNode("plan", "Placement and sequence planning", ["Allocate rows/columns and segments", "Validate geometry and inventory", "Generate true construction order"], "service"),
                DiagramNode("status", "Planner outcome", ["VALIDATED or review-ready status", "INFEASIBLE/INCOMPLETE diagnostics", "Inventory, shortages, steps, HTML", "Teacher approval only if approvable"], "decision"),
            ]),
        ],
        [
            DiagramEdge("intent", "guard", "validate"), DiagramEdge("guard", "rodin", "constrained prompt"),
            DiagramEdge("rodin", "bang", "approved current OBJ"), DiagramEdge("bang", "voxel", "segmented geometry"),
            DiagramEdge("voxel", "contacts", "physical regions"), DiagramEdge("contacts", "renders", "tables + mappings"),
            DiagramEdge("voxel", "catalog", "voxel plan"), DiagramEdge("contacts", "catalog", "functional interfaces"),
            DiagramEdge("catalog", "plan", "allowed parts + ledger"), DiagramEdge("plan", "status", "validated result"),
            DiagramEdge("status", "renders", "true steps / diagnostics", True), DiagramEdge("status", "guard", "bounded regeneration guidance", True),
        ],
    ),
    Diagram(
        "document-generation",
        "Lesson Bundle Composition, Validation, and Publication",
        "How one approved source of truth becomes three audience-specific documents without changing build facts.",
        [
            DiagramSection("Approved source package", "Immutable facts for this generation", [
                DiagramNode("doccontext", "Document context", ["Story analysis and teacher planning", "Framework/RAG evidence", "Approved object and movement map", "Validated inventory and true build plan"], "artifact"),
                DiagramNode("images", "Validated visual assets", ["Final build reference", "Segment and multiview images", "True step isometric and placement sheets", "No placeholder/demo art"], "artifact"),
            ]),
            DiagramSection("Audience-specific composition", "Shared facts, different instructional purpose", [
                DiagramNode("lesson", "Teacher lesson-plan agent", ["Read / Learn & Explore / Invent", "Objectives, timing, prompts, standards", "Final-build and multiview overview"], "service"),
                DiagramNode("activity", "Student activity-guide agent", ["Concise classroom directions", "Vocabulary, phonics, real-world links", "Reflection and final-build overview"], "service"),
                DiagramNode("slides", "Slide-companion agent", ["Classroom-facing sequence", "Kid-friendly prompts and vocabulary", "All validated build steps and images"], "service"),
            ]),
            DiagramSection("Deterministic quality gates", "Per-document validation before publication", [
                DiagramNode("schema", "Structure and content checks", ["Required sections and minimum depth", "Audience-appropriate tone", "Framework anchors where expected", "No placeholders or missing values"], "decision"),
                DiagramNode("visual", "Image and build checks", ["Files exist and decode", "Notebook/build images included", "Step count agrees with build plan", "No stale generation references"], "decision"),
                DiagramNode("approval", "Teacher document review", ["Preview each document", "Refine independently", "All three must be valid and approved"], "actor"),
            ]),
            DiagramSection("Publication outputs", "Download and maintenance artifacts", [
                DiagramNode("sourcesout", "Editable/debug sources", ["Markdown content", "Structured JSON context", "Validation report"], "artifact"),
                DiagramNode("pdfs", "Classroom PDFs", ["lesson_plan.pdf", "activity_guide.pdf", "slide_companion.pdf"], "artifact"),
                DiagramNode("download", "FastAPI download responses", ["Stable filenames and media type", "Streams bytes, never local paths", "Future: persist to GCS with lifecycle"], "system"),
            ]),
        ],
        [
            DiagramEdge("doccontext", "lesson", "compose"), DiagramEdge("doccontext", "activity", "compose"),
            DiagramEdge("doccontext", "slides", "compose"), DiagramEdge("images", "lesson", "selected visuals"),
            DiagramEdge("images", "activity", "final overview"), DiagramEdge("images", "slides", "all build visuals"),
            DiagramEdge("lesson", "schema", "draft"), DiagramEdge("activity", "schema", "draft"),
            DiagramEdge("slides", "schema", "draft"), DiagramEdge("schema", "visual", "content valid"),
            DiagramEdge("visual", "approval", "preview"), DiagramEdge("approval", "sourcesout", "approve"),
            DiagramEdge("approval", "pdfs", "publish"), DiagramEdge("sourcesout", "download", "serve"),
            DiagramEdge("pdfs", "download", "serve"),
        ],
    ),
    Diagram(
        "automatic-recovery",
        "Bounded Automatic Simplification and Exception Recovery",
        "How the happy path self-corrects block and segment excess before asking the teacher to regenerate the model.",
        [
            DiagramSection("Measure candidate", "Every Bang/notebook generation", [
                DiagramNode("candidate", "Current candidate", ["Approved OBJ generation", "Bang source segments", "Voxelized physical regions", "Initial block approximation"], "artifact"),
                DiagramNode("measure", "Constraint evaluator", ["Block count vs max_validated_blocks", "Source/physical segments vs max_semantic_parts", "Required/moving-part preservation", "Inventory and planner status"], "decision"),
            ]),
            DiagramSection("Bounded internal recovery", "No teacher action while safe options remain", [
                DiagramNode("tune", "Attempt policy", ["Try configured voxel-size candidates", "Consolidate adjacent static fragments", "Protect moving and required regions", "Never invent unsupported catalog parts"], "service"),
                DiagramNode("score", "Attempt scoring", ["Feasibility first", "Then semantic preservation", "Then block/segment margin", "Record diagnostics and selected attempt"], "decision"),
                DiagramNode("retry", "Bounded retry controller", ["Small fixed attempt budget", "No repeated paid Rodin calls", "May rerun local physicalization/planner", "Stops on valid or no improvement"], "service"),
            ]),
            DiagramSection("Outcome and user path", "Only actionable states reach the teacher", [
                DiagramNode("valid", "Approachable result", ["Validated or CSP review-ready", "Inventory feasible", "True steps and images available", "Teacher can inspect and approve"], "artifact"),
                DiagramNode("recoverable", "Local exception result", ["Best candidate retained", "Clear blocking measurements", "Label/connector refinement if sufficient"], "artifact"),
                DiagramNode("regen", "Model regeneration required", ["Specific Rodin prompt update", "Prefilled build constraints", "Old segment/build generations invalidated", "One prominent return action"], "artifact"),
            ]),
        ],
        [
            DiagramEdge("candidate", "measure", "evaluate"), DiagramEdge("measure", "valid", "all gates pass"),
            DiagramEdge("measure", "tune", "safe recovery exists"), DiagramEdge("tune", "score", "candidate attempt"),
            DiagramEdge("score", "retry", "rank"), DiagramEdge("retry", "measure", "next local attempt", True),
            DiagramEdge("score", "recoverable", "best local result"), DiagramEdge("recoverable", "regen", "still blocked"),
            DiagramEdge("regen", "candidate", "new model generation", True),
        ],
    ),
    Diagram(
        "security-operations",
        "Security, Observability, and Recovery Architecture",
        "Cross-cutting controls for identity, secrets, untrusted content, monitoring, ephemeral state, and operator recovery.",
        [
            DiagramSection("Identity and trust", "Who may call what", [
                DiagramNode("human", "Human access", ["Research-team GCP IAM", "Optional future institutional sign-in", "Separate admin ingestion from classroom use"], "actor"),
                DiagramNode("sa", "Cloud Run service account", ["Vertex AI User", "Cloud SQL Client", "GCS object access", "Secret accessor and log writer"], "system"),
                DiagramNode("third", "Third-party boundary", ["Hyper3D credential only", "No student PII in 3D prompts", "Review processing and retention terms"], "external"),
            ]),
            DiagramSection("Application controls", "Prevent unsafe state or content transitions", [
                DiagramNode("filesec", "File and path controls", ["MIME/extension and size checks", "Safe filenames and isolated job dirs", "Known-artifact download mapping"], "service"),
                DiagramNode("promptsec", "Prompt and evidence controls", ["Retrieved text treated as untrusted", "Structured output validation", "Deterministic readiness and stage gates"], "service"),
                DiagramNode("secretsec", "Secret handling", ["Secret Manager references", "No secrets in Git, logs, screenshots", "ADC preferred over key files"], "service"),
            ]),
            DiagramSection("Telemetry and operations", "Diagnose without leaking content", [
                DiagramNode("logs", "Structured logs", ["request/session correlation", "stage, provider, duration, fallback", "auto-recovery and planner status", "exception class; no story/prompt bodies"], "external"),
                DiagramNode("alerts", "Operational signals", ["5xx and startup failures", "DB/retrieval degradation", "Long model/segment jobs", "Planner incomplete and download errors"], "external"),
                DiagramNode("health", "Health and readiness", ["Process liveness", "Model and database configuration", "Retrieval fallback mode", "Catalog/runtime availability"], "service"),
            ]),
            DiagramSection("Backup and recovery", "Current posture and hardening path", [
                DiagramNode("durable", "Durable data", ["Cloud SQL backups and PITR policy", "GCS versioning/lifecycle as configured", "Corpus re-ingestion is reproducible"], "datastore"),
                DiagramNode("ephemeral", "Ephemeral state risk", ["In-memory sessions", "Container-local 3D and PDFs", "Lost when Cloud Run instance is replaced"], "datastore"),
                DiagramNode("runbook", "Operator recovery", ["Check readiness and logs", "Rollback Cloud Run revision", "Restore/reload corpus", "Future: durable session/artifact store"], "artifact"),
            ]),
        ],
        [
            DiagramEdge("human", "sa", "deploys/configures"), DiagramEdge("sa", "third", "authorized API calls"),
            DiagramEdge("sa", "secretsec", "runtime identity"), DiagramEdge("filesec", "logs", "safe events"),
            DiagramEdge("promptsec", "logs", "validation outcomes"), DiagramEdge("secretsec", "logs", "access metadata only"),
            DiagramEdge("logs", "alerts", "metrics and patterns"), DiagramEdge("health", "alerts", "readiness failures"),
            DiagramEdge("durable", "runbook", "restore"), DiagramEdge("ephemeral", "runbook", "recreate / warn"),
        ],
    ),
]


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/aptos.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
    ]
    if bold:
        candidates = [Path("C:/Windows/Fonts/aptos-bold.ttf"), Path("C:/Windows/Fonts/arialbd.ttf")] + candidates
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def _xml_escape(value: str) -> str:
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


NODE_STYLES = {
    "actor": ("#FFF7E3", GOLD, NAVY),
    "system": (SKY, TEAL, NAVY),
    "service": ("#F6F8FA", "#7A9E9F", INK),
    "datastore": ("#F1ECFA", "#7C5AA6", NAVY),
    "external": ("#FFF0ED", CORAL, NAVY),
    "decision": ("#FFF8D8", "#B88413", NAVY),
    "artifact": ("#ECF7EE", GREEN, NAVY),
}


def _wrapped_lines(value: str, width: int) -> list[str]:
    return textwrap.wrap(value, width=max(18, width), break_long_words=False, break_on_hyphens=False) or [value]


def _edge_anchor(source: tuple[int, int, int, int], target: tuple[int, int, int, int]) -> tuple[tuple[int, int], tuple[int, int]]:
    sx1, sy1, sx2, sy2 = source
    tx1, ty1, tx2, ty2 = target
    scx, scy = (sx1 + sx2) // 2, (sy1 + sy2) // 2
    tcx, tcy = (tx1 + tx2) // 2, (ty1 + ty2) // 2
    if abs(tcx - scx) >= abs(tcy - scy):
        if tcx >= scx:
            return (sx2, scy), (tx1, tcy)
        return (sx1, scy), (tx2, tcy)
    if tcy >= scy:
        return (scx, sy2), (tcx, ty1)
    return (scx, sy1), (tcx, ty2)


def export_diagram(diagram: Diagram) -> None:
    width = 2000
    section_count = len(diagram.sections)
    margin_x = 42
    gap_x = 24
    section_w = (width - margin_x * 2 - gap_x * (section_count - 1)) // section_count
    node_gap = 22
    node_top = 270
    node_specs: dict[str, tuple[DiagramNode, tuple[int, int, int, int], int]] = {}
    section_heights: list[int] = []
    for section_index, section in enumerate(diagram.sections):
        x = margin_x + section_index * (section_w + gap_x)
        y = node_top
        for node in section.nodes:
            wrap_width = 43 if section_count <= 3 else 31
            line_count = sum(len(_wrapped_lines(item, wrap_width)) for item in node.items)
            node_h = max(142, 72 + line_count * 27 + len(node.items) * 7)
            node_specs[node.key] = (node, (x + 18, y, x + section_w - 18, y + node_h), section_index)
            y += node_h + node_gap
        section_heights.append(y - node_gap + 24)
    diagram_bottom = max(section_heights, default=700)
    height = diagram_bottom + 176
    png = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(png)
    draw.rectangle((0, 0, width, 104), fill=NAVY)
    draw.text((54, 26), diagram.title, font=_font(38, True), fill="white")
    draw.text((54, 126), diagram.purpose, font=_font(22), fill=INK)
    svg_parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="white"/>',
        f'<rect width="{width}" height="104" fill="{NAVY}"/>',
        f'<text x="54" y="65" font-family="Arial" font-size="38" font-weight="700" fill="white">{_xml_escape(diagram.title)}</text>',
        f'<text x="54" y="148" font-family="Arial" font-size="22" fill="{INK}">{_xml_escape(diagram.purpose)}</text>',
        '<defs><marker id="arrow" viewBox="0 0 10 10" refX="9" refY="5" markerWidth="7" markerHeight="7" orient="auto-start-reverse"><path d="M 0 0 L 10 5 L 0 10 z" fill="#667085"/></marker></defs>',
    ]
    section_bottom = diagram_bottom
    for section_index, section in enumerate(diagram.sections):
        x = margin_x + section_index * (section_w + gap_x)
        boundary = (x, 180, x + section_w, section_bottom)
        draw.rounded_rectangle(boundary, radius=16, fill="#FBFCFD", outline="#B8C2CC", width=2)
        draw.rectangle((x, 180, x + section_w, 228), fill="#E9EEF3")
        draw.text((x + 18, 190), section.title, font=_font(24, True), fill=NAVY)
        draw.text((x + 18, 232), section.subtitle, font=_font(15), fill=MUTED)
        svg_parts.extend([
            f'<rect x="{x}" y="180" width="{section_w}" height="{section_bottom-180}" rx="16" fill="#FBFCFD" stroke="#B8C2CC" stroke-width="2"/>',
            f'<rect x="{x}" y="180" width="{section_w}" height="48" fill="#E9EEF3"/>',
            f'<text x="{x+18}" y="213" font-family="Arial" font-size="24" font-weight="700" fill="{NAVY}">{_xml_escape(section.title)}</text>',
            f'<text x="{x+18}" y="248" font-family="Arial" font-size="15" fill="{MUTED}">{_xml_escape(section.subtitle)}</text>',
        ])

    for edge in diagram.edges:
        if edge.source not in node_specs or edge.target not in node_specs:
            continue
        _, source_box, _ = node_specs[edge.source]
        _, target_box, _ = node_specs[edge.target]
        (x1, y1), (x2, y2) = _edge_anchor(source_box, target_box)
        dash = (12, 8) if edge.dashed else None
        if dash:
            total = max(1, int(((x2 - x1) ** 2 + (y2 - y1) ** 2) ** 0.5))
            for offset in range(0, total, sum(dash)):
                start = offset / total
                end = min(1.0, (offset + dash[0]) / total)
                draw.line((x1 + (x2-x1)*start, y1 + (y2-y1)*start, x1 + (x2-x1)*end, y1 + (y2-y1)*end), fill=MUTED, width=3)
        else:
            draw.line((x1, y1, x2, y2), fill=MUTED, width=3)
        angle_x, angle_y = x2 - x1, y2 - y1
        mag = max(1.0, (angle_x * angle_x + angle_y * angle_y) ** 0.5)
        ux, uy = angle_x / mag, angle_y / mag
        left = (x2 - ux * 18 - uy * 8, y2 - uy * 18 + ux * 8)
        right = (x2 - ux * 18 + uy * 8, y2 - uy * 18 - ux * 8)
        draw.polygon([(x2, y2), left, right], fill=MUTED)
        midx, midy = (x1 + x2) // 2, (y1 + y2) // 2
        label_lines = _wrapped_lines(edge.label, 24)
        label_w = min(250, max(118, max(len(line) for line in label_lines) * 8 + 24))
        label_h = 24 + len(label_lines) * 17
        draw.rounded_rectangle((midx - label_w//2, midy - label_h//2, midx + label_w//2, midy + label_h//2), radius=7, fill="white", outline="#D0D5DD")
        for line_no, line in enumerate(label_lines):
            draw.text((midx - label_w//2 + 12, midy - label_h//2 + 9 + line_no*17), line, font=_font(13), fill=MUTED)
        dash_attr = ' stroke-dasharray="12 8"' if edge.dashed else ""
        svg_parts.extend([
            f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{MUTED}" stroke-width="3"{dash_attr} marker-end="url(#arrow)"/>',
            f'<rect x="{midx-label_w//2}" y="{midy-label_h//2}" width="{label_w}" height="{label_h}" rx="7" fill="white" stroke="#D0D5DD"/>',
        ])
        for line_no, line in enumerate(label_lines):
            svg_parts.append(f'<text x="{midx}" y="{midy-label_h//2+22+line_no*17}" text-anchor="middle" font-family="Arial" font-size="13" fill="{MUTED}">{_xml_escape(line)}</text>')

    for node, box, _ in node_specs.values():
        x1, y1, x2, y2 = box
        fill, border, text_color = NODE_STYLES.get(node.kind, NODE_STYLES["service"])
        draw.rounded_rectangle(box, radius=12, fill=fill, outline=border, width=3)
        draw.rectangle((x1, y1, x2, y1 + 48), fill=fill)
        draw.text((x1 + 16, y1 + 11), node.title, font=_font(22, True), fill=text_color)
        kind_label = node.kind.upper()
        badge_w = max(64, len(kind_label) * 9 + 18)
        draw.rounded_rectangle((x2 - badge_w - 12, y1 + 10, x2 - 12, y1 + 38), radius=9, fill="white", outline=border)
        draw.text((x2 - badge_w, y1 + 16), kind_label, font=_font(11, True), fill=border)
        svg_parts.extend([
            f'<rect x="{x1}" y="{y1}" width="{x2-x1}" height="{y2-y1}" rx="12" fill="{fill}" stroke="{border}" stroke-width="3"/>',
            f'<text x="{x1+16}" y="{y1+32}" font-family="Arial" font-size="22" font-weight="700" fill="{text_color}">{_xml_escape(node.title)}</text>',
            f'<rect x="{x2-badge_w-12}" y="{y1+10}" width="{badge_w}" height="28" rx="9" fill="white" stroke="{border}"/>',
            f'<text x="{x2-badge_w//2-12}" y="{y1+29}" text-anchor="middle" font-family="Arial" font-size="11" font-weight="700" fill="{border}">{kind_label}</text>',
        ])
        yy = y1 + 61
        wrap_width = 43 if section_count <= 3 else 31
        for item in node.items:
            wrapped = _wrapped_lines(item, wrap_width)
            draw.ellipse((x1 + 17, yy + 7, x1 + 25, yy + 15), fill=border)
            svg_parts.append(f'<circle cx="{x1+21}" cy="{yy+11}" r="4" fill="{border}"/>')
            for line_no, line in enumerate(wrapped):
                draw.text((x1 + 34, yy + line_no * 27), line, font=_font(18), fill=INK)
                svg_parts.append(f'<text x="{x1+34}" y="{yy+18+line_no*27}" font-family="Arial" font-size="18" fill="{INK}">{_xml_escape(line)}</text>')
            yy += len(wrapped) * 27 + 7

    legend_y = section_bottom + 36
    draw.text((54, legend_y), "Legend", font=_font(18, True), fill=NAVY)
    legend_x = 150
    for kind in ["actor", "system", "service", "datastore", "external", "decision", "artifact"]:
        fill, border, _ = NODE_STYLES[kind]
        draw.rounded_rectangle((legend_x, legend_y - 2, legend_x + 24, legend_y + 22), radius=5, fill=fill, outline=border, width=2)
        draw.text((legend_x + 32, legend_y), kind.title(), font=_font(14), fill=MUTED)
        svg_parts.extend([
            f'<rect x="{legend_x}" y="{legend_y-2}" width="24" height="24" rx="5" fill="{fill}" stroke="{border}" stroke-width="2"/>',
            f'<text x="{legend_x+32}" y="{legend_y+17}" font-family="Arial" font-size="14" fill="{MUTED}">{kind.title()}</text>',
        ])
        legend_x += 210
    draw.text((54, legend_y + 52), "Solid arrow: primary runtime/data flow   |   Dashed arrow: feedback, invalidation, fallback, or recovery path", font=_font(15), fill=MUTED)
    svg_parts.append(f'<text x="54" y="{legend_y+69}" font-family="Arial" font-size="15" fill="{MUTED}">Solid arrow: primary runtime/data flow | Dashed arrow: feedback, invalidation, fallback, or recovery path</text>')

    png.save(IMAGES / f"diagram-{diagram.slug}.png", optimize=True)
    svg_parts.append("</svg>")
    (IMAGES / f"diagram-{diagram.slug}.svg").write_text("\n".join(svg_parts), encoding="utf-8")
    mermaid = ["flowchart LR", f"  %% {diagram.purpose}"]
    for section_index, section in enumerate(diagram.sections):
        mermaid.append(f'  subgraph S{section_index}["{section.title} - {section.subtitle}"]')
        mermaid.append("    direction TB")
        for node in section.nodes:
            label = f"{node.title}<br/>" + "<br/>".join(node.items)
            mermaid.append(f'    {node.key}["{label}"]:::{node.kind}')
        mermaid.append("  end")
    for edge in diagram.edges:
        operator = "-.->" if edge.dashed else "-->"
        mermaid.append(f"  {edge.source} {operator}|{edge.label}| {edge.target}")
    for kind, (fill, border, color) in NODE_STYLES.items():
        mermaid.append(f"  classDef {kind} fill:{fill},stroke:{border},color:{color},stroke-width:2px;")
    (IMAGES / f"diagram-{diagram.slug}.mmd").write_text("\n".join(mermaid) + "\n", encoding="utf-8")


def copy_reference_visuals() -> None:
    source = ROOT / "work" / "build_jobs" / "session_9bed4a65-06d9-4eb1-9eac-4a21abfb1ecc"
    copies = {
        source / "notebook_outputs" / "segment_visualization.png": "08-segment-review.png",
        source / "notebook_outputs" / "brick_approximation.png": "09-validated-build.png",
        source / "notebook_outputs" / "notebook_step_03.png": "10-build-step.png",
        source / "notebook_outputs" / "notebook_step_03_multiview.png": "11-build-step-multiview.png",
    }
    for src, name in copies.items():
        if src.exists():
            shutil.copy2(src, IMAGES / name)


def render_bundle_contact_sheet() -> None:
    source = ROOT / "work" / "build_jobs" / "session_9bed4a65-06d9-4eb1-9eac-4a21abfb1ecc" / "lesson_bundle"
    thumbs: list[tuple[str, Image.Image]] = []
    try:
        import fitz
    except ImportError:
        return
    for label, filename in [
        ("Teacher Lesson Plan", "lesson_plan.pdf"),
        ("Student Activity Guide", "activity_guide.pdf"),
        ("Slide Companion", "slide_companion.pdf"),
    ]:
        path = source / filename
        if not path.exists():
            continue
        doc = fitz.open(path)
        page = doc.load_page(0)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.3, 1.3), alpha=False)
        image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        image.thumbnail((430, 560))
        thumbs.append((label, image.copy()))
        doc.close()
    if not thumbs:
        return
    canvas = Image.new("RGB", (1540, 720), "white")
    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, 0, 1540, 82), fill=NAVY)
    draw.text((42, 23), "Validated Lesson Bundle", font=_font(30, True), fill="white")
    x = 58
    for label, image in thumbs:
        draw.rounded_rectangle((x - 10, 112, x + 450, 676), radius=10, fill=LIGHT, outline="#D0D5DD", width=2)
        draw.text((x, 126), label, font=_font(20, True), fill=NAVY)
        canvas.paste(image, (x, 168))
        x += 500
    canvas.save(IMAGES / "12-lesson-bundle.png", optimize=True)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def _set_cell_width(cell, dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def _style_doc(doc: Document, title: str, preset: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    header = section.header.paragraphs[0]
    header.text = "KIDSPARK AI / BRICKSMART  |  FINAL HANDOFF"
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string("667085")
    _add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5 if preset == "compact" else 9.8)
    normal.font.color.rgb = RGBColor.from_string("243142")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT if preset == "compact" else WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size, color, before, after in [
        ("Title", 26, "17324D", 0, 12),
        ("Heading 1", 17, "17324D", 16, 8),
        ("Heading 2", 14, "2F7F8D", 12, 6),
        ("Heading 3", 11.5, "17324D", 9, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Code Block" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(8)
        code_style.paragraph_format.left_indent = Inches(0.18)
        code_style.paragraph_format.right_indent = Inches(0.18)
        code_style.paragraph_format.space_after = Pt(6)


def _strip_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text


def _docx_rich_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            p.add_run(text[pos:match.start()])
        token = match.group(0)
        run = p.add_run(token.strip("*`"))
        if token.startswith("**"):
            run.bold = True
        elif token.startswith("*"):
            run.italic = True
        else:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string("A23E34")
        pos = match.end()
    if pos < len(text):
        p.add_run(text[pos:])
    return p


def _parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
            continue
        rows.append(cells)
    return rows


def markdown_to_docx(source: Path, destination: Path, preset: str) -> None:
    doc = Document()
    _style_doc(doc, source.stem, preset)
    lines = source.read_text(encoding="utf-8").splitlines()
    in_code = False
    code: list[str] = []
    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph("\n".join(code), style="Code Block")
                _shade_cell if False else None
                code = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code.append(line)
            i += 1
            continue
        if line.strip() == "\\pagebreak":
            doc.add_page_break()
            i += 1
            continue
        image_match = re.match(r"!\[([^]]*)\]\(([^)]+)\)", line.strip())
        if image_match:
            path = (source.parent / image_match.group(2)).resolve()
            if path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(path), width=Inches(6.55))
                cap = doc.add_paragraph(image_match.group(1))
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(8)
                cap.runs[0].font.color.rgb = RGBColor.from_string("667085")
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            block = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            rows = _parse_table(block)
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                width = 9360 // cols
                for r, row in enumerate(rows):
                    for c in range(cols):
                        cell = table.cell(r, c)
                        cell.text = _strip_inline(row[c] if c < len(row) else "")
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        _set_cell_width(cell, width)
                        if r == 0:
                            _shade_cell(cell, "E8F4F7")
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                                run.font.color.rgb = RGBColor.from_string("17324D")
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title" if first_title else "Heading 1")
            p.add_run(_strip_inline(line[2:]))
            first_title = False
        elif line.startswith("## "):
            doc.add_paragraph(_strip_inline(line[3:]), style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(_strip_inline(line[4:]), style="Heading 2")
        elif line.startswith("#### "):
            doc.add_paragraph(_strip_inline(line[5:]), style="Heading 3")
        elif re.match(r"^[-*] ", line):
            p = _docx_rich_paragraph(doc, line[2:], style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.38)
            p.paragraph_format.first_line_indent = Inches(-0.18)
        elif re.match(r"^\d+\. ", line):
            number_match = re.match(r"^(\d+)\. (.*)$", line)
            number = number_match.group(1) if number_match else "1"
            body = number_match.group(2) if number_match else line
            p = _docx_rich_paragraph(doc, f"{number}.  {body}")
            p.paragraph_format.left_indent = Inches(0.38)
            p.paragraph_format.first_line_indent = Inches(-0.18)
        elif line.startswith("> "):
            p = _docx_rich_paragraph(doc, line[2:])
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.color.rgb = RGBColor.from_string("17324D")
                run.italic = True
        elif line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p_pr = p._p.get_or_add_pPr()
            border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:color"), "2F7F8D")
            border.append(bottom)
            p_pr.append(border)
        elif line.strip():
            _docx_rich_paragraph(doc, line)
        i += 1
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def _rl_styles(preset: str):
    base = getSampleStyleSheet()
    base.add(ParagraphStyle(name="KSBody", fontName="Helvetica", fontSize=9.2, leading=12.1, textColor=colors.HexColor(INK), alignment=TA_JUSTIFY if preset == "narrative" else TA_LEFT, spaceAfter=6))
    base.add(ParagraphStyle(name="KSTitle", fontName="Helvetica-Bold", fontSize=24, leading=28, textColor=colors.HexColor(NAVY), alignment=TA_CENTER, spaceAfter=18))
    base.add(ParagraphStyle(name="KSH1", fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor(NAVY), spaceBefore=12, spaceAfter=7))
    base.add(ParagraphStyle(name="KSH2", fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor(TEAL), spaceBefore=9, spaceAfter=5))
    base.add(ParagraphStyle(name="KSH3", fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=colors.HexColor(NAVY), spaceBefore=7, spaceAfter=4))
    base.add(ParagraphStyle(name="KSCode", fontName="Courier", fontSize=6.8, leading=8.5, backColor=colors.HexColor("#F2F4F7"), leftIndent=8, rightIndent=8, borderPadding=5, spaceAfter=7))
    base.add(ParagraphStyle(name="KSCaption", fontName="Helvetica-Oblique", fontSize=7.2, leading=9, textColor=colors.HexColor(MUTED), alignment=TA_CENTER, spaceAfter=8))
    base.add(ParagraphStyle(name="KSQuote", fontName="Helvetica-Oblique", fontSize=9, leading=12, textColor=colors.HexColor(NAVY), leftIndent=18, rightIndent=18, backColor=colors.HexColor(SKY), borderPadding=8, spaceAfter=8))
    base.add(ParagraphStyle(name="KSNumbered", parent=base["KSBody"], leftIndent=18, firstLineIndent=-18))
    return base


def markdown_to_pdf(source: Path, destination: Path, preset: str) -> None:
    styles = _rl_styles(preset)
    story = []
    lines = source.read_text(encoding="utf-8").splitlines()
    in_code = False
    code: list[str] = []
    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                story.append(Paragraph(_xml_escape("\n".join(code)).replace("\n", "<br/>"), styles["KSCode"]))
                code = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code.append(line)
            i += 1
            continue
        if line.strip() == "\\pagebreak":
            story.append(PageBreak())
            i += 1
            continue
        image_match = re.match(r"!\[([^]]*)\]\(([^)]+)\)", line.strip())
        if image_match:
            path = (source.parent / image_match.group(2)).resolve()
            if path.exists():
                with Image.open(path) as im:
                    w, h = im.size
                max_w, max_h = 6.85 * inch, 8.0 * inch
                scale = min(max_w / w, max_h / h)
                story.append(RLImage(str(path), width=w * scale, height=h * scale))
                story.append(Paragraph(_xml_escape(image_match.group(1)), styles["KSCaption"]))
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            block = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            rows = _parse_table(block)
            if rows:
                data = [[Paragraph(_xml_escape(_strip_inline(c)), styles["KSBody"]) for c in row] for row in rows]
                table = Table(data, repeatRows=1, hAlign="LEFT", colWidths=[6.8 * inch / len(data[0])] * len(data[0]))
                table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(SKY)),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(NAVY)),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D5DD")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                story.append(table)
                story.append(Spacer(1, 7))
            continue
        text = _xml_escape(_strip_inline(line.lstrip("#>0123456789.- ")))
        if line.startswith("# "):
            if not first_title:
                story.append(PageBreak())
            story.append(Paragraph(text, styles["KSTitle"]))
            first_title = False
        elif line.startswith("## "):
            story.append(Paragraph(text, styles["KSH1"]))
        elif line.startswith("### "):
            story.append(Paragraph(text, styles["KSH2"]))
        elif line.startswith("#### "):
            story.append(Paragraph(text, styles["KSH3"]))
        elif re.match(r"^[-*] ", line):
            story.append(ListFlowable([ListItem(Paragraph(_xml_escape(_strip_inline(line[2:])), styles["KSBody"]))], bulletType="bullet", leftIndent=18, bulletFontSize=6))
        elif re.match(r"^\d+\. ", line):
            number_match = re.match(r"^(\d+)\. (.*)$", line)
            number = number_match.group(1) if number_match else "1"
            body = number_match.group(2) if number_match else line
            numbered = Paragraph(
                f"<b>{number}.</b>&nbsp;&nbsp;{_xml_escape(_strip_inline(body))}",
                styles["KSNumbered"],
            )
            story.append(numbered)
        elif line.startswith("> "):
            story.append(Paragraph(_xml_escape(_strip_inline(line[2:])), styles["KSQuote"]))
        elif line.strip() == "---":
            story.append(Spacer(1, 8))
        elif line.strip():
            story.append(Paragraph(_xml_escape(_strip_inline(line)), styles["KSBody"]))
        i += 1

    def header_footer(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(colors.HexColor(NAVY))
        canvas.rect(0, LETTER[1] - 28, LETTER[0], 28, fill=1, stroke=0)
        canvas.setFont("Helvetica-Bold", 7.5)
        canvas.setFillColor(colors.white)
        canvas.drawString(0.72 * inch, LETTER[1] - 19, "KIDSPARK AI / BRICKSMART  |  FINAL HANDOFF")
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor(MUTED))
        canvas.drawRightString(LETTER[0] - 0.72 * inch, 0.34 * inch, f"Page {doc.page}")
        canvas.restoreState()

    destination.parent.mkdir(parents=True, exist_ok=True)
    pdf = SimpleDocTemplate(
        str(destination),
        pagesize=LETTER,
        leftMargin=0.72 * inch,
        rightMargin=0.72 * inch,
        topMargin=0.68 * inch,
        bottomMargin=0.55 * inch,
        title=source.stem.replace("_", " "),
        author="KidSpark AI / BrickSmart Project Team",
    )
    pdf.build(story, onFirstPage=header_footer, onLaterPages=header_footer)


def export_openapi() -> None:
    os.environ.setdefault("KIDSPARK_OFFLINE_MODE", "true")
    os.environ.setdefault("DATABASE_REQUIRED", "false")
    backend = ROOT / "backend"
    sys.path.insert(0, str(backend))
    try:
        from api.main import app
    except ModuleNotFoundError as exc:
        if exc.name != "fastapi":
            raise
        app_python = shutil.which("python")
        if not app_python or Path(app_python).resolve() == Path(sys.executable).resolve():
            raise RuntimeError("Application Python with FastAPI is required to export OpenAPI") from exc
        env = os.environ.copy()
        env.setdefault("KIDSPARK_OFFLINE_MODE", "true")
        env.setdefault("DATABASE_REQUIRED", "false")
        command = (
            "import json,sys; "
            f"sys.path.insert(0, {str(backend)!r}); "
            "from api.main import app; "
            "schema=app.openapi(); "
            "schema['info']['description']='Sanitized OpenAPI snapshot for the KidSpark AI application API.'; "
            f"open({str(OPENAPI / 'kidspark-api.openapi.json')!r}, 'w', encoding='utf-8').write(json.dumps(schema, indent=2))"
        )
        subprocess.run([app_python, "-c", command], check=True, env=env, cwd=ROOT)
        return

    schema = app.openapi()
    schema["info"]["description"] = "Sanitized OpenAPI snapshot for the KidSpark AI application API."
    OPENAPI.mkdir(parents=True, exist_ok=True)
    (OPENAPI / "kidspark-api.openapi.json").write_text(json.dumps(schema, indent=2), encoding="utf-8")


def main() -> int:
    technical_only = "--technical-only" in sys.argv[1:]
    unknown_args = [arg for arg in sys.argv[1:] if arg != "--technical-only"]
    if unknown_args:
        raise SystemExit(f"Unknown arguments: {', '.join(unknown_args)}")
    HANDOFF.mkdir(parents=True, exist_ok=True)
    IMAGES.mkdir(parents=True, exist_ok=True)
    OPENAPI.mkdir(parents=True, exist_ok=True)
    for diagram in DIAGRAMS:
        export_diagram(diagram)
    copy_reference_visuals()
    render_bundle_contact_sheet()
    export_openapi()
    pairs = [
        (DOCS / "KIDSPARK_TECHNICAL_DESIGN.md", HANDOFF / "KidSpark_Technical_Design.docx", HANDOFF / "KidSpark_Technical_Design.pdf", "compact"),
    ]
    if not technical_only:
        pairs.append(
            (DOCS / "KIDSPARK_PROJECT_OVERVIEW.md", HANDOFF / "KidSpark_Project_Overview.docx", HANDOFF / "KidSpark_Project_Overview.pdf", "narrative")
        )
    for source, docx, pdf, preset in pairs:
        if not source.exists():
            raise FileNotFoundError(source)
        markdown_to_docx(source, docx, preset)
        markdown_to_pdf(source, pdf, preset)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
